"""
Office文档处理器 - 支持Word和Excel文件的文本提取和OCR识别
"""

import os
import json
import time
import pickle
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from loguru import logger

# 尝试导入Office文档处理库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx库未安装，Word文档处理功能将不可用")

try:
    import openpyxl
    from openpyxl import load_workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logger.warning("openpyxl库未安装，Excel文档处理功能将不可用")

try:
    import xlrd
    XLS_AVAILABLE = True
except ImportError:
    XLS_AVAILABLE = False
    logger.warning("xlrd库未安装，旧版Excel文档处理功能将不可用")

from .config import OUTPUT_DIR, PICKLES_DIR, PROMPTS, REMOTE_OCR_CONFIG
from .ocr_engine import OCREngine
from .llm_processor import LLMProcessor
from .remote_ocr_client import RemoteOCRClient


class OfficeProcessor:
    """Office文档处理器，支持Word和Excel文档的文本提取和OCR识别"""
    
    def __init__(self, use_gpu: bool = True, ocr_engine=None):
        """
        初始化Office文档处理器
        
        Args:
            use_gpu: 是否使用GPU
            ocr_engine: 外部OCR引擎实例（可选）
        """
        if ocr_engine:
            self.ocr_engine = ocr_engine
        else:
            self.ocr_engine = OCREngine(use_gpu=use_gpu)
        
        # 初始化LLM处理器
        self.llm_processor = LLMProcessor()
        self.mode = "快速"
        
        # 检查依赖库可用性
        self.docx_available = DOCX_AVAILABLE
        self.xlsx_available = XLSX_AVAILABLE
        self.xls_available = XLS_AVAILABLE
        
        # 远程OCR客户端
        self.remote_client = None
        if REMOTE_OCR_CONFIG.get("enabled", False):
            try:
                self.remote_client = RemoteOCRClient(REMOTE_OCR_CONFIG["server_url"])
                if self.remote_client.check_server_health():
                    logger.info("远程GPU Office服务连接成功")
                else:
                    logger.warning("远程GPU Office服务连接失败，将使用本地处理")
                    self.remote_client = None
            except Exception as e:
                logger.warning(f"远程Office客户端初始化失败: {e}")
                self.remote_client = None
        
        logger.info(f"Office处理器初始化完成 - Word: {self.docx_available}, Excel: {self.xlsx_available}, 旧Excel: {self.xls_available}")
    
    def set_mode(self, mode: str):
        """设置解析模式：快速/精细"""
        if mode not in ("快速", "精细"):
            return
        self.mode = mode
        logger.info(f"Office处理器模式设置为: {self.mode}")
    
    def process_office_document(self, file_path: str, output_name: str = None) -> Dict:
        """
        处理Office文档文件
        
        Args:
            file_path: 文档文件路径
            output_name: 输出名称，如果为None则使用文件名
            
        Returns:
            处理结果字典
        """
        if not Path(file_path).exists():
            return {'status': 'error', 'message': f'文档文件不存在: {file_path}'}
        
        if output_name is None:
            output_name = Path(file_path).stem
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            # 优先使用远程OCR
            if self.remote_client:
                try:
                    logger.info(f"使用远程GPU Office处理文档: {output_name}")
                    result = self.remote_client.process_office(file_path, f"{output_name}{file_ext}")
                    if result.get('status') == 'success':
                        return result
                    else:
                        logger.warning(f"远程Office处理失败，回退到本地处理: {result.get('message', '未知错误')}")
                        if not REMOTE_OCR_CONFIG.get("fallback_to_local", True):
                            return result
                except Exception as e:
                    logger.warning(f"远程Office处理异常，回退到本地处理: {e}")
                    if not REMOTE_OCR_CONFIG.get("fallback_to_local", True):
                        return {'status': 'error', 'message': str(e)}
            
            # 本地处理
            # 创建输出目录
            output_path = OUTPUT_DIR / output_name
            output_path.mkdir(exist_ok=True)
            
            # 根据文件类型进行相应处理
            if file_ext == '.docx':
                result = self._process_docx_file(file_path, output_path, output_name)
            elif file_ext == '.doc':
                result = self._process_doc_file(file_path, output_path, output_name)
            elif file_ext == '.xlsx':
                result = self._process_xlsx_file(file_path, output_path, output_name)
            elif file_ext == '.xls':
                result = self._process_xls_file(file_path, output_path, output_name)
            else:
                return {'status': 'error', 'message': f'不支持的文件类型: {file_ext}'}
            
            # 清理临时文件
            self._cleanup_temp_files(output_path)
            
            return result
            
        except Exception as e:
            logger.error(f"Office文档处理失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def _process_docx_file(self, file_path: str, output_path: Path, output_name: str) -> Dict:
        """处理Word文档(.docx)"""
        if not self.docx_available:
            return {'status': 'error', 'message': 'python-docx库未安装，无法处理Word文档'}
        
        try:
            logger.info(f"开始处理Word文档: {Path(file_path).name}")
            
            # 打开Word文档
            doc = Document(file_path)
            
            # 提取文本内容
            all_texts = []
            paragraphs = []
            
            # 提取段落文本
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
                    all_texts.append({
                        'type': 'paragraph',
                        'index': i + 1,
                        'text': text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # 提取表格文本
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                        if cell_text:
                            all_texts.append({
                                'type': 'table_cell',
                                'table_index': table_idx + 1,
                                'row_index': row_idx + 1,
                                'column_index': cell_idx + 1,
                                'text': cell_text
                            })
                    table_data.append(row_data)
                tables.append({
                    'table_index': table_idx + 1,
                    'data': table_data
                })
            
            # 合并所有文本
            full_text = '\n'.join(paragraphs)
            
            # 生成摘要和关键词
            summary_result = self._generate_summary(full_text)
            
            # 保存结果
            result = {
                'status': 'success',
                'output_path': str(output_path),
                'file_type': 'docx',
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables),
                'texts': all_texts,
                'paragraphs': paragraphs,
                'tables': tables,
                'summary': summary_result.get('summary', ''),
                'keywords': summary_result.get('keywords', []),
                'hybrid_summary': summary_result.get('hybrid_summary', ''),
                'markdown_content': summary_result.get('markdown_content', ''),
                'part_summaries': summary_result.get('part_summaries', []),
                'categories': summary_result.get('categories', []),
                'category_descriptions': summary_result.get('category_descriptions', {}),
                'category_confidence': summary_result.get('category_confidence', 0.0),
                'tags': summary_result.get('tags', [])
            }
            
            # 保存到pickle文件
            self._save_to_pickle(result, output_name)
            
            logger.info(f"Word文档处理完成: {Path(file_path).name}, 共{len(paragraphs)}段, {len(tables)}表")
            return result
            
        except Exception as e:
            logger.error(f"处理Word文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def _process_doc_file(self, file_path: str, output_path: Path, output_name: str) -> Dict:
        """处理旧版Word文档(.doc) - 需要转换为图片进行OCR"""
        try:
            logger.info(f"开始处理旧版Word文档: {Path(file_path).name}")
            
            # 对于.doc文件，我们需要先转换为图片，然后进行OCR
            # 这里使用一个简化的方法，实际项目中可能需要更复杂的转换
            
            # 尝试使用python-docx2txt或其他库
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                if text:
                    # 如果成功提取文本，按段落分割
                    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                    full_text = '\n'.join(paragraphs)
                    
                    all_texts = []
                    for i, paragraph in enumerate(paragraphs):
                        all_texts.append({
                            'type': 'paragraph',
                            'index': i + 1,
                            'text': paragraph
                        })
                    
                    # 生成摘要和关键词
                    summary_result = self._generate_summary(full_text)
                    
                    result = {
                        'status': 'success',
                        'output_path': str(output_path),
                        'file_type': 'doc',
                        'total_paragraphs': len(paragraphs),
                        'total_tables': 0,
                        'texts': all_texts,
                        'paragraphs': paragraphs,
                        'tables': [],
                        'summary': summary_result.get('summary', ''),
                        'keywords': summary_result.get('keywords', []),
                        'hybrid_summary': summary_result.get('hybrid_summary', ''),
                        'markdown_content': summary_result.get('markdown_content', ''),
                        'part_summaries': summary_result.get('part_summaries', []),
                        'categories': summary_result.get('categories', []),
                        'category_descriptions': summary_result.get('category_descriptions', {}),
                        'category_confidence': summary_result.get('category_confidence', 0.0),
                        'tags': summary_result.get('tags', [])
                    }
                    
                    self._save_to_pickle(result, output_name)
                    logger.info(f"旧版Word文档处理完成: {Path(file_path).name}")
                    return result
                    
            except ImportError:
                logger.warning("docx2txt库未安装，无法处理.doc文件")
                return {'status': 'error', 'message': 'docx2txt库未安装，无法处理.doc文件'}
            
            return {'status': 'error', 'message': '无法提取.doc文件内容'}
            
        except Exception as e:
            logger.error(f"处理旧版Word文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def _process_xlsx_file(self, file_path: str, output_path: Path, output_name: str) -> Dict:
        """处理Excel文档(.xlsx)"""
        if not self.xlsx_available:
            return {'status': 'error', 'message': 'openpyxl库未安装，无法处理Excel文档'}
        
        try:
            logger.info(f"开始处理Excel文档: {Path(file_path).name}")
            
            # 打开Excel工作簿
            workbook = load_workbook(file_path, data_only=True)
            
            all_texts = []
            sheets_data = []
            
            # 处理每个工作表
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = {
                    'sheet_name': sheet_name,
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column,
                    'data': []
                }
                
                # 提取单元格数据
                for row_idx, row in enumerate(sheet.iter_rows(), 1):
                    row_data = []
                    for col_idx, cell in enumerate(row, 1):
                        cell_value = str(cell.value) if cell.value is not None else ""
                        row_data.append(cell_value)
                        
                        if cell_value.strip():
                            all_texts.append({
                                'type': 'cell',
                                'sheet_name': sheet_name,
                                'row_index': row_idx,
                                'column_index': col_idx,
                                'text': cell_value,
                                'cell_reference': f"{sheet_name}!{cell.coordinate}"
                            })
                    
                    if any(cell.strip() for cell in row_data):  # 只保存非空行
                        sheet_data['data'].append(row_data)
                
                sheets_data.append(sheet_data)
            
            # 合并所有文本内容
            all_cell_texts = [text['text'] for text in all_texts if text['text'].strip()]
            full_text = '\n'.join(all_cell_texts)
            
            # 生成摘要和关键词
            summary_result = self._generate_summary(full_text)
            
            # 保存结果
            result = {
                'status': 'success',
                'output_path': str(output_path),
                'file_type': 'xlsx',
                'total_sheets': len(sheets_data),
                'total_cells': len(all_texts),
                'texts': all_texts,
                'sheets_data': sheets_data,
                'summary': summary_result.get('summary', ''),
                'keywords': summary_result.get('keywords', []),
                'hybrid_summary': summary_result.get('hybrid_summary', ''),
                'markdown_content': summary_result.get('markdown_content', ''),
                'part_summaries': summary_result.get('part_summaries', []),
                'categories': summary_result.get('categories', []),
                'category_descriptions': summary_result.get('category_descriptions', {}),
                'category_confidence': summary_result.get('category_confidence', 0.0),
                'tags': summary_result.get('tags', [])
            }
            
            # 保存到pickle文件
            self._save_to_pickle(result, output_name)
            
            logger.info(f"Excel文档处理完成: {Path(file_path).name}, 共{len(sheets_data)}表, {len(all_texts)}单元格")
            return result
            
        except Exception as e:
            logger.error(f"处理Excel文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def _process_xls_file(self, file_path: str, output_path: Path, output_name: str) -> Dict:
        """处理旧版Excel文档(.xls)"""
        if not self.xls_available:
            return {'status': 'error', 'message': 'xlrd库未安装，无法处理旧版Excel文档'}
        
        try:
            logger.info(f"开始处理旧版Excel文档: {Path(file_path).name}")
            
            # 打开Excel工作簿
            workbook = xlrd.open_workbook(file_path)
            
            all_texts = []
            sheets_data = []
            
            # 处理每个工作表
            for sheet_idx in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sheet_idx)
                sheet_name = sheet.name
                
                sheet_data = {
                    'sheet_name': sheet_name,
                    'max_row': sheet.nrows,
                    'max_column': sheet.ncols,
                    'data': []
                }
                
                # 提取单元格数据
                for row_idx in range(sheet.nrows):
                    row_data = []
                    for col_idx in range(sheet.ncols):
                        cell_value = str(sheet.cell_value(row_idx, col_idx))
                        row_data.append(cell_value)
                        
                        if cell_value.strip():
                            all_texts.append({
                                'type': 'cell',
                                'sheet_name': sheet_name,
                                'row_index': row_idx + 1,
                                'column_index': col_idx + 1,
                                'text': cell_value,
                                'cell_reference': f"{sheet_name}!{xlrd.colname(col_idx)}{row_idx + 1}"
                            })
                    
                    if any(cell.strip() for cell in row_data):  # 只保存非空行
                        sheet_data['data'].append(row_data)
                
                sheets_data.append(sheet_data)
            
            # 合并所有文本内容
            all_cell_texts = [text['text'] for text in all_texts if text['text'].strip()]
            full_text = '\n'.join(all_cell_texts)
            
            # 生成摘要和关键词
            summary_result = self._generate_summary(full_text)
            
            # 保存结果
            result = {
                'status': 'success',
                'output_path': str(output_path),
                'file_type': 'xls',
                'total_sheets': len(sheets_data),
                'total_cells': len(all_texts),
                'texts': all_texts,
                'sheets_data': sheets_data,
                'summary': summary_result.get('summary', ''),
                'keywords': summary_result.get('keywords', []),
                'hybrid_summary': summary_result.get('hybrid_summary', ''),
                'markdown_content': summary_result.get('markdown_content', ''),
                'part_summaries': summary_result.get('part_summaries', []),
                'categories': summary_result.get('categories', []),
                'category_descriptions': summary_result.get('category_descriptions', {}),
                'category_confidence': summary_result.get('category_confidence', 0.0),
                'tags': summary_result.get('tags', [])
            }
            
            # 保存到pickle文件
            self._save_to_pickle(result, output_name)
            
            logger.info(f"旧版Excel文档处理完成: {Path(file_path).name}, 共{len(sheets_data)}表, {len(all_texts)}单元格")
            return result
            
        except Exception as e:
            logger.error(f"处理旧版Excel文档失败: {e}")
            return {'status': 'error', 'message': str(e)}
    
    def _generate_summary(self, text: str) -> Dict:
        """生成内容摘要"""
        try:
            if not text.strip():
                return {
                    'summary': '文档内容为空',
                    'keywords': [],
                    'hybrid_summary': '无内容可分析',
                    'markdown_content': '无内容可转换',
                    'part_summaries': [],
                    'categories': [],
                    'category_descriptions': {},
                    'category_confidence': 0.0,
                    'tags': []
                }
            
            # 使用LLM生成各种类型的内容
            summary = self.llm_processor.generate_summary(text)
            keywords = self.llm_processor.extract_keywords(text)
            hybrid_summary = self.llm_processor.generate_hybrid_summary(text)
            markdown_content = self.llm_processor.convert_to_markdown(text)
            
            # 分类与打标签
            try:
                categories = self.llm_processor.classify(text)
                tags = self.llm_processor.generate_tags(text)
            except Exception as e:
                logger.warning(f"分类/打标签阶段异常: {e}")
                categories = {"categories": [], "confidence": 0.0}
                tags = []
            
            # 生成段落摘要
            part_summaries = []
            try:
                # 将文本分段处理
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                for i, paragraph in enumerate(paragraphs[:5]):  # 只处理前5段
                    if len(paragraph) > 100:  # 只对较长的段落生成摘要
                        part_summary = self.llm_processor.generate_summary(paragraph)
                        part_summaries.append({
                            'paragraph_index': i + 1,
                            'original_length': len(paragraph),
                            'summary': part_summary
                        })
            except Exception as e:
                logger.warning(f"生成段落摘要失败: {e}")
            
            return {
                'summary': summary,
                'keywords': keywords,
                'hybrid_summary': hybrid_summary,
                'markdown_content': markdown_content,
                'part_summaries': part_summaries,
                'categories': [c.get('name') for c in categories.get('categories', [])],
                'category_descriptions': {c.get('name'): c.get('description') for c in categories.get('categories', []) if c.get('name')},
                'category_confidence': categories.get('confidence', 0.0),
                'tags': tags
            }
            
        except Exception as e:
            logger.error(f"生成摘要失败: {e}")
            return {
                'summary': '摘要生成失败',
                'keywords': [],
                'hybrid_summary': '分析失败',
                'markdown_content': '转换失败',
                'part_summaries': [],
                'categories': [],
                'category_descriptions': {},
                'category_confidence': 0.0,
                'tags': []
            }
    
    def _save_to_pickle(self, result: Dict, output_name: str):
        """保存结果到pickle文件"""
        try:
            pickle_path = PICKLES_DIR / output_name
            pickle_path.mkdir(exist_ok=True)
            
            # 保存主要结果
            with open(pickle_path / "result.pkl", 'wb') as f:
                pickle.dump(result, f)
            
            # 保存文本向量（如果需要）
            if result['texts']:
                texts = [text['text'] for text in result['texts'] if isinstance(text, dict) and 'text' in text]
                with open(pickle_path / "texts.pkl", 'wb') as f:
                    pickle.dump(texts, f)
            
            logger.info(f"结果已保存到: {pickle_path}")
            
        except Exception as e:
            logger.error(f"保存pickle文件失败: {e}")
    
    def _cleanup_temp_files(self, output_path: Path):
        """清理临时文件"""
        try:
            # 删除临时文件
            for temp_file in output_path.glob("*.tmp"):
                temp_file.unlink()
            
            logger.info("临时文件清理完成")
            
        except Exception as e:
            logger.error(f"清理临时文件失败: {e}")
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        formats = []
        if self.docx_available:
            formats.append('.docx')
        if self.xlsx_available:
            formats.append('.xlsx')
        if self.xls_available:
            formats.append('.xls')
        # .doc 需要额外库支持
        formats.append('.doc')
        return formats
